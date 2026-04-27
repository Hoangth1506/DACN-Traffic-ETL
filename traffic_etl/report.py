from __future__ import annotations

from pathlib import Path

import pandas as pd
from docx import Document
from docx.shared import Inches
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfbase import pdfmetrics
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle


def build_reports(
    out_dir: Path,
    observations: pd.DataFrame,
    fusion: pd.DataFrame,
    stats: pd.DataFrame,
    tests: pd.DataFrame,
    chart_paths: dict[str, Path],
    metadata: dict,
) -> tuple[Path, Path]:
    out_dir.mkdir(parents=True, exist_ok=True)
    docx_path = out_dir / "non_iid_etl_report.docx"
    pdf_path = out_dir / "non_iid_etl_report.pdf"
    _build_docx(docx_path, observations, fusion, stats, tests, chart_paths, metadata)
    _build_pdf(pdf_path, fusion, stats, tests, metadata)
    return docx_path, pdf_path


def _build_docx(path: Path, observations, fusion, stats, tests, chart_paths, metadata) -> None:
    doc = Document()
    doc.add_heading("ETL hop nhat du lieu giao thong Non-IID tai 3 node bien", 0)
    doc.add_paragraph(
        "Bao cao nay trinh bay pipeline Extract - Transform - Load cho 3 khu vuc "
        "Ly Thuong Kiet, Cong Hoa va Truong Chinh. API key duoc doc tu bien moi truong "
        "TOMTOM_API_KEY va khong duoc ghi vao ma nguon hay bao cao."
    )
    doc.add_paragraph(metadata["data_readiness_note"])
    doc.add_heading("Nguon goc du lieu", 1)
    for item in metadata["lineage"]:
        doc.add_paragraph(item, style="List Bullet")
    doc.add_paragraph(f"Anh minh chung node bien: {metadata['edge_node_image']}")

    doc.add_heading("How to transform", 1)
    for item in _transform_steps():
        doc.add_paragraph(item, style="List Number")

    doc.add_heading("Ket qua hop nhat node", 1)
    _add_table(doc, fusion)
    doc.add_heading("Minh chung Non-IID", 1)
    doc.add_paragraph(
        "Non-IID duoc chung minh bang viec phan phoi velocity khac nhau giua cac node "
        "va ban chat nguon khac nhau giua TomTom live flow, OSM topology, PyTorch YOLO video "
        "va fallback metadata."
    )
    _add_table(doc, stats)
    _add_table(doc, tests)
    doc.add_heading("Bieu do xuat kem", 1)
    for name, chart in chart_paths.items():
        doc.add_paragraph(f"{name}: {chart}")
    doc.add_heading("Cua so thu thap", 1)
    doc.add_paragraph(
        f"Ban dau: {metadata['initial_history_days']} ngay, cac khung {metadata['windows']}; "
        f"tu dong tren GitHub Actions trong {metadata['auto_collection_months']} thang."
    )
    doc.save(path)


def _build_pdf(path: Path, fusion, stats, tests, metadata) -> None:
    _register_font()
    styles = getSampleStyleSheet()
    font_name = "DejaVuSans" if "DejaVuSans" in pdfmetrics.getRegisteredFontNames() else "Helvetica"
    for style in styles.byName.values():
        style.fontName = font_name
    doc = SimpleDocTemplate(str(path), pagesize=A4, rightMargin=1.5 * cm, leftMargin=1.5 * cm)
    story = [
        Paragraph("ETL hop nhat du lieu giao thong Non-IID tai 3 node bien", styles["Title"]),
        Spacer(1, 0.3 * cm),
        Paragraph("Nguon du lieu: toa do thu cong trong config/nodes.yaml, TomTom Traffic Flow, OSM Overpass, anh 3 diem.", styles["BodyText"]),
        Spacer(1, 0.2 * cm),
        Paragraph("How to transform: chuan hoa schema, tinh congestion ratio, LOS, density proxy, gan lineage, hop nhat co trong so.", styles["BodyText"]),
        Spacer(1, 0.3 * cm),
        Paragraph("Ket qua hop nhat node", styles["Heading2"]),
        _pdf_table(fusion),
        Spacer(1, 0.3 * cm),
        Paragraph("Thong ke Non-IID", styles["Heading2"]),
        _pdf_table(stats),
        Spacer(1, 0.3 * cm),
        Paragraph("Kiem dinh phan phoi", styles["Heading2"]),
        _pdf_table(tests),
        Spacer(1, 0.2 * cm),
        Paragraph(f"Trang thai du lieu: {metadata['data_readiness_note']}", styles["BodyText"]),
        Paragraph(f"Cua so thu thap: {metadata['windows']}; GitHub auto collection: {metadata['auto_collection_months']} thang.", styles["BodyText"]),
    ]
    doc.build(story)


def _add_table(doc: Document, df: pd.DataFrame) -> None:
    if df.empty:
        doc.add_paragraph("Khong co du lieu.")
        return
    display = df.copy()
    table = doc.add_table(rows=1, cols=len(display.columns))
    table.style = "Table Grid"
    for idx, col in enumerate(display.columns):
        table.rows[0].cells[idx].text = str(col)
    for _, row in display.iterrows():
        cells = table.add_row().cells
        for idx, col in enumerate(display.columns):
            cells[idx].text = str(row[col])


def _pdf_table(df: pd.DataFrame) -> Table:
    if df.empty:
        data = [["empty"]]
    else:
        view = df.head(8).copy()
        data = [list(view.columns)] + [[str(v)[:28] for v in row] for row in view.values.tolist()]
    table = Table(data, repeatRows=1)
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.lightgrey),
                ("GRID", (0, 0), (-1, -1), 0.25, colors.grey),
                ("FONT", (0, 0), (-1, -1), "DejaVuSans" if "DejaVuSans" in pdfmetrics.getRegisteredFontNames() else "Helvetica", 7),
            ]
        )
    )
    return table


def _register_font() -> None:
    candidates = [
        Path("C:/Windows/Fonts/DejaVuSans.ttf"),
        Path("C:/Windows/Fonts/arial.ttf"),
    ]
    for candidate in candidates:
        if candidate.exists():
            try:
                pdfmetrics.registerFont(TTFont("DejaVuSans", str(candidate)))
            except Exception:
                pass
            return


def _transform_steps() -> list[str]:
    return [
        "Extract: lay toa do node thu cong tu config/nodes.yaml; lay live traffic flow tai cac diem mau quanh node neu co TOMTOM_API_KEY; lay road graph tu OSM Overpass.",
        "Transform schema: dua moi observation ve node_id, timestamp, lat/lon, velocity_kmph, free_flow_kmph, confidence, source_name, source_api.",
        "Transform metric: congestion_ratio = 1 - currentSpeed/freeFlowSpeed; density_proxy = congestion_ratio * 30; LOS theo nguong van toc A-F trong paper.",
        "Transform lineage: giu extracted_at, raw_path, source_api, source_name de minh chung nguon goc tung record.",
        "Video AI: PyTorch YOLO nhan dien motorcycle/car/bus/truck, tracking centroid de tinh van toc va line-crossing de tinh luu luong.",
        "Load: xuat raw JSON, processed CSV/JSON/Parquet neu co engine, node fusion va bao cao DOCX/PDF.",
        "Fusion: w = alpha*confidence + beta*recency + gamma*source_quality; chuan hoa w trong tung node roi tinh trung binh co trong so.",
    ]
